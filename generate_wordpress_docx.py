"""
STEP 4A:
Generate DOCX files but send to OpenAI to be checkd by grammar.
"""
import os
import json
import html2text
from docx import Document
from anyio import Path
from openai import OpenAI
import time
import asyncio
from dotenv import load_dotenv

load_dotenv()

     
api_key = os.getenv("OPENAI_API_KEY")
 
def grammar_teacher(text):
    try:
        # Construct a prompt for grammar correction
        # print (text)
        # print ('1223')
        prompt = f"You are a copywriter and an English teacher. Correct my text but do not rewrite, just improve and correct it.: \n\n'{text}'"

        client = OpenAI(
            # This is the default and can be omitted
            api_key=os.environ.get("OPENAI_API_KEY"),
        )
        # Sending a request to the GPT-4 Turbo model
        response = client.chat.completions.create(
                messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="gpt-4-turbo-preview",
            )
        
        print (response)
        return response.choices[0].message.content.strip()
    except Exception as e:
        return str(e)
      
def create_word_document(article_data, image_dir):
    doc = Document()
    doc.add_heading(article_data['title'], level=1)
     # Example for content, in a real scenario you would need to properly handle HTML content and encoding
    
    h = html2text.HTML2Text()
    h.ignore_links = True
    raw_text = h.handle(article_data['body'])

    corrected_text = grammar_teacher(raw_text)
    doc.add_paragraph(corrected_text)
    # Here you could add logic to include images from image_dir
    return doc


async def main():
    article_dir = 'articles/'
    doc_output_dir = 'docx/'
    await Path(doc_output_dir).mkdir(exist_ok=True) 

    article_files = [f for f in os.listdir(article_dir) if f.endswith('.json')] # [:1]   Limit to first 2 articles

    for filename in article_files:
        with open(os.path.join(article_dir, filename), 'r') as file:
            article_data = json.load(file)
            article_id = article_data['id']
            article_title = article_data['title'].replace(' ', '_').replace('/', '_')
            # Generate a unique filename for each article's Word document
            doc_filename = f'{article_title}_{article_id}.docx'
            # Check if the Word document already exists
            if os.path.exists(os.path.join(doc_output_dir, doc_filename)):
                print(f"Word document for article {article_id} already exists at {doc_output_dir}{doc_filename}. Skipping.")
            else:
                article_image_dir = os.path.join(article_dir, str(article_id))
                doc = create_word_document(article_data, article_image_dir)
                doc.save(os.path.join(doc_output_dir, doc_filename))
                print(f"Generated Word document for article {article_id} at {doc_output_dir}{doc_filename}")

if __name__ == '__main__':
    asyncio.run(main())
